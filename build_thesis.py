"""
build_thesis.py
==============
Ensambla la tesis en .docx con DISEÑO ACADÉMICO PROFESIONAL:
- Cuerpo Times New Roman 12, justificado, interlineado 1,5, con guionado automático.
- Encabezados con tratamiento tipográfico (color y filete) y control de viudas/huérfanas.
- Tablas estilo *booktabs* (solo filetes horizontales: superior, de cabecera e inferior).
- Encabezado de página con título corto + filete; pie con número de página.
- Márgenes de encuadernación (izq. 3 cm) y numeración romana/arábiga por secciones.
- Portada depurada con filetes; índices general, de tablas y de figuras (campos Word).

Salida: outputs/Tesis_Francisco_Rietta.docx
"""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
DOCS = ROOT / "docs"
OUT = ROOT / "outputs" / "Tesis_Francisco_Rietta.docx"

TITULO = ("Impacto de las variables macroeconómicas en los retornos accionarios de las "
          "empresas de cobre con exposición a Chile: una comparación entre el mercado "
          "bursátil internacional y el chileno, 2004–2024")
TITULO_CORTO = "Macroeconomía y retornos del sector cobre en Chile"

NAVY = RGBColor(0x00, 0x00, 0x00)   # negro (norma APA para títulos)
RULE = "808080"
GRAY = RGBColor(0x59, 0x59, 0x59)
BODY_FONT = "Georgia"               # tipografía USS / APA
HEAD_FONT = "Georgia"               # misma fuente en títulos (APA)

# Estructura USS (Señor de Sipán) / esquema de informe de investigación
CAPS = ["uss_01_introduccion.md", "uss_02_metodo.md", "uss_03_resultados.md",
        "uss_04_conclusiones.md", "referencias.md", "anexos.md"]

ABREV = [
    ("APT", "Arbitrage Pricing Theory (teoría de valoración por arbitraje)"),
    ("ADF", "Augmented Dickey-Fuller (prueba de raíz unitaria)"),
    ("ARDL", "Autoregressive Distributed Lag (rezagos distribuidos)"),
    ("BCCh", "Banco Central de Chile"),
    ("CD", "Cross-section Dependence (prueba de Pesaran)"),
    ("CIPS", "Cross-sectionally augmented IPS (raíz unitaria en panel)"),
    ("CMF", "Comisión para el Mercado Financiero"),
    ("EMBI", "Emerging Markets Bond Index (riesgo soberano)"),
    ("FEVD", "Forecast Error Variance Decomposition"),
    ("FE", "Fixed Effects (efectos fijos)"),
    ("FRED", "Federal Reserve Economic Data"),
    ("IMACEC", "Índice Mensual de Actividad Económica"),
    ("IPSA", "Índice de Precios Selectivo de Acciones"),
    ("KPSS", "Kwiatkowski-Phillips-Schmidt-Shin (prueba de estacionariedad)"),
    ("TPM", "Tasa de Política Monetaria"),
    ("VAR", "Vector Autoregression (vector autorregresivo)"),
    ("VECM", "Vector Error Correction Model"),
    ("VIX", "CBOE Volatility Index"),
]


# ---------- estilos base ----------
def set_base_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT; st.font.size = Pt(12)
    st.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE   # interlineado doble (norma de tesis)
    pf.space_after = Pt(8); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.widow_control = True
    # Niveles de título según APA 7: N1 centrado negrita; N2 izq. negrita; N3 izq. negrita cursiva
    apa = [("Heading 1", 14, WD_ALIGN_PARAGRAPH.CENTER, False),
           ("Heading 2", 12, WD_ALIGN_PARAGRAPH.LEFT, False),
           ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT, True)]
    for h, sz, al, ital in apa:
        s = doc.styles[h]
        s.font.name = HEAD_FONT; s.font.size = Pt(sz); s.font.bold = True; s.font.italic = ital
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.alignment = al
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
    # estilo de leyenda (Caption)
    try:
        cap = doc.styles["Caption"]
        cap.font.name = BODY_FONT; cap.font.size = Pt(10); cap.font.italic = True
        cap.font.color.rgb = GRAY
    except KeyError:
        pass


def enable_hyphenation(doc):
    el = doc.settings.element
    auto = OxmlElement("w:autoHyphenation"); auto.set(qn("w:val"), "true"); el.append(auto)
    lim = OxmlElement("w:consecutiveHyphenLimit"); lim.set(qn("w:val"), "3"); el.append(lim)
    # zona de guionado estrecha (~0,25 cm) => justificación más pareja, menos huecos
    zone = OxmlElement("w:hyphenationZone"); zone.set(qn("w:val"), "142"); el.append(zone)
    nocap = OxmlElement("w:doNotHyphenateCaps"); nocap.set(qn("w:val"), "true"); el.append(nocap)


def _fld(run, instr, default="1"):
    for t, attr in [("begin", None), (None, instr), ("separate", None),
                    ("text", default), ("end", None)]:
        if attr is not None:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = attr
            run._r.append(it)
        elif t == "text":
            tt = OxmlElement("w:t"); tt.text = default; run._r.append(tt)
        else:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), t); run._r.append(fc)


def set_pgnum_format(section, fmt, start=None):
    pg = section._sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); section._sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def footer_pagenum(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.text = ""
    run = p.add_run(); _fld(run, "PAGE"); run.font.size = Pt(10); run.font.name = BODY_FONT


def running_header(section, text=None):
    # APA: número de página arriba a la derecha; sin running head en formato estudiante.
    section.header.is_linked_to_previous = False
    p = section.header.paragraphs[0]; p.text = ""; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(); _fld(run, "PAGE"); run.font.size = Pt(11); run.font.name = BODY_FONT


def field_block(doc, instr, placeholder):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    run = p.add_run(); _fld(run, instr, placeholder)


def add_runs_with_bold(paragraph, text):
    # negrita **..**, código `..`, matemática inline $..$, itálica *..*, marca [COMPLETAR..]
    pat = r"(\*\*.+?\*\*|`[^`]+`|\$[^$\n]+\$|\*[^*\n]+?\*|\[COMPLETAR[^\]]*\])"
    for part in re.split(pat, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10.5)
        elif part.startswith("$") and part.endswith("$"):
            paragraph.add_run(part[1:-1]).italic = True   # notación matemática
        elif part.startswith("[COMPLETAR"):
            r = paragraph.add_run(part); r.italic = True; r.font.color.rgb = GRAY
        elif len(part) > 2 and part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def clean_heading(text):
    return re.sub(r"\s*\((borrador|texto)[^)]*\)", "", text).strip()


# ---------- leyendas con campo SEQ ----------
def caption(doc, label, text, above=False):
    try:
        p = doc.add_paragraph(style="Caption")
    except KeyError:
        p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8 if not above else 2)
    p.paragraph_format.space_after = Pt(2 if above else 10)
    r = p.add_run(f"{label} "); r.bold = True; r.font.size = Pt(10); r.italic = False
    fs = OxmlElement("w:fldSimple"); fs.set(qn("w:instr"), f" SEQ {label} \\* ARABIC ")
    rr = OxmlElement("w:r"); tt = OxmlElement("w:t"); tt.text = "1"; rr.append(tt); fs.append(rr)
    p._p.append(fs)
    if text:
        r2 = p.add_run(f". {text}"); r2.font.size = Pt(10); r2.italic = True


# ---------- tablas estilo booktabs ----------
def _cell_border(cell, edge, sz, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
    el = borders.find(qn("w:" + edge))
    if el is None:
        el = OxmlElement("w:" + edge); borders.append(el)
    for k, v in (("w:val", "single"), ("w:sz", str(sz)), ("w:space", "0"), ("w:color", color)):
        el.set(qn(k), v)


def _cell_pad(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, w in (("top", 60), ("bottom", 60), ("left", 80), ("right", 80)):
        e = OxmlElement("w:" + edge); e.set(qn("w:w"), str(w)); e.set(qn("w:type"), "dxa"); mar.append(e)
    tcPr.append(mar)


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not cells:
        return
    caption(doc, "Tabla", "", above=True)
    tbl = doc.add_table(rows=len(cells), cols=len(cells[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    nrows = len(cells)
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j >= len(tbl.rows[i].cells):
                continue
            c = tbl.rows[i].cells[j]; cp = c.paragraphs[0]; cp.text = ""
            cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.line_spacing = 1.0
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            add_runs_with_bold(cp, val.replace("**", "") if i == 0 else val)
            for run in cp.runs:
                run.font.size = Pt(10); run.font.name = BODY_FONT
                if i == 0:
                    run.bold = True
            _cell_pad(c)
            # booktabs: filete superior (cabecera), inferior de cabecera, inferior final
            if i == 0:
                _cell_border(c, "top", 14, RULE)
                _cell_border(c, "bottom", 8, RULE)
            if i == nrows - 1:
                _cell_border(c, "bottom", 14, RULE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, cap_text, path):
    img = ROOT / path
    if not img.exists():
        return
    doc.add_picture(str(img), width=Inches(5.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(8)
    caption(doc, "Figura", cap_text)


def unwrap_paragraphs(md):
    """Une las líneas de texto envueltas en un único párrafo lógico, incluidas las
    continuaciones de ítems de lista (mejor flujo; evita que negritas o frases se corten
    en los saltos de línea del markdown)."""
    # Línea tras la cual NO se debe anexar (bloque no extensible o vacía):
    def is_break(s):
        return (not s.strip()) or bool(re.match(r"\s*(#{1,6}\s|\||!\[|>|---|\$\$)", s))
    # Línea de continuación: texto plano que no inicia un nuevo elemento estructural:
    def is_cont(s):
        return bool(s.strip()) and not re.match(r"\s*(#{1,6}\s|[-*]\s|\d+\.\s|\||!\[|>|---|\$\$)", s)
    out = []
    for raw in md.split("\n"):
        s = raw.rstrip()
        if out and is_cont(s) and not is_break(out[-1]):
            out[-1] += " " + s.strip()
        else:
            out.append(s)
    return "\n".join(out)


def render_markdown(doc, md_text, hanging=False):
    lines = unwrap_paragraphs(md_text).split("\n"); i = 0; tablebuf = []
    def flush():
        nonlocal tablebuf
        if tablebuf:
            add_table(doc, tablebuf); tablebuf = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("|") and "|" in line[1:]:
            tablebuf.append(line); i += 1; continue
        flush()
        if not line.strip() or line.startswith(">") or line.startswith("---"):
            i += 1; continue
        m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if m:
            add_figure(doc, m.group(1), m.group(2)); i += 1; continue
        meq = re.match(r"^\s*\$\$(.+?)\$\$\s*$", line)
        if meq:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
            r = p.add_run(meq.group(1).strip()); r.italic = True
            i += 1; continue
        if line.startswith("# "):
            doc.add_heading(clean_heading(line[2:]), 1); i += 1; continue
        if line.startswith("## "):
            doc.add_heading(clean_heading(line[3:]), 2); i += 1; continue
        if line.startswith("### "):
            doc.add_heading(clean_heading(line[4:]), 3); i += 1; continue
        if line.startswith("#### "):
            doc.add_heading(clean_heading(line[5:]), 3); i += 1; continue
        if re.match(r"^\s*[-*] ", line):
            add_runs_with_bold(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*] ", "", line)); i += 1; continue
        mn = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if mn:
            # numeración LITERAL (preserva 1,2,3 por lista; evita la continuación global de Word)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.8)
            add_runs_with_bold(p, f"{mn.group(1)}. {mn.group(2)}")
            i += 1; continue
        p = doc.add_paragraph()
        if hanging:  # sangría francesa para referencias APA
            p.paragraph_format.left_indent = Cm(1.27)
            p.paragraph_format.first_line_indent = Cm(-1.27)
        add_runs_with_bold(p, line); i += 1
    flush()


# ---------- preliminares ----------
def _no_hyphen(p):
    pPr = p._p.get_or_add_pPr()
    sup = OxmlElement("w:suppressAutoHyphens"); sup.set(qn("w:val"), "true"); pPr.append(sup)


def centered(doc, text, size, bold=False, italic=False, space=6, color=None, font=BODY_FONT):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space); p.paragraph_format.line_spacing = 1.15
    _no_hyphen(p)
    r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.name = font
    if color:
        r.font.color.rgb = color
    return p


def hrule(doc, color=RULE, sz=12, space=10):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space); p.paragraph_format.space_after = Pt(space)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); pPr.append(pbdr)
    b = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", str(sz)), ("w:space", "1"), ("w:color", color)):
        b.set(qn(k), v)
    pbdr.append(b)


def add_cover(doc):
    BLACK = RGBColor(0, 0, 0)
    centered(doc, "UNIVERSIDAD SEÑOR DE SIPÁN", 16, bold=True, space=2, color=BLACK)
    centered(doc, "FACULTAD DE [Facultad]", 13, bold=True, space=2, color=BLACK)
    centered(doc, "ESCUELA PROFESIONAL DE [Escuela / Programa]", 13, bold=True, space=10, color=BLACK)
    logo = next((c for c in (ROOT/"assets"/"logo_uss.png", ROOT.parent/"tesis-plataforma"/"logo_uss.png")
                 if c.exists()), None)
    if logo:
        doc.add_picture(str(logo), width=Inches(1.7)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for _ in range(2):
            doc.add_paragraph()
    centered(doc, "TESIS", 14, bold=True, space=12, color=BLACK)
    centered(doc, TITULO.upper(), 14, bold=True, space=12, color=BLACK)
    centered(doc, "PARA OPTAR EL TÍTULO PROFESIONAL / GRADO ACADÉMICO DE [Grado o Título]",
             12, space=24, color=BLACK)
    centered(doc, "Autor(es):", 12, bold=True, space=0, color=BLACK)
    centered(doc, "Rietta Francisco", 12, space=10, color=BLACK)
    centered(doc, "Asesor:", 12, bold=True, space=0, color=BLACK)
    centered(doc, "[Grado académico. Apellidos y Nombres del asesor]", 12, space=10, color=BLACK)
    centered(doc, "Línea de Investigación:", 12, bold=True, space=0, color=BLACK)
    centered(doc, "[Línea de Investigación aprobada por la Escuela]", 12, space=24, color=BLACK)
    centered(doc, "Pimentel – Perú", 12, bold=True, space=2, color=BLACK)
    centered(doc, "2026", 12, bold=True, color=BLACK)


def page(doc): doc.add_page_break()


def add_approval(doc):
    doc.add_heading("Página de aprobación", 1)
    p = doc.add_paragraph(); add_runs_with_bold(p, "La presente tesis ha sido revisada y aprobada por "
        "la comisión examinadora abajo firmante, como requisito para optar el título profesional / grado "
        "académico de [Grado o Título] en la Universidad Señor de Sipán.")
    for rol in ["Profesor guía", "Profesor informante", "Director(a) del programa"]:
        doc.add_paragraph()
        l = doc.add_paragraph("__________________________________"); l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        l.paragraph_format.space_after = Pt(0)
        n = doc.add_paragraph(rol); n.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in n.runs:
            r.font.color.rgb = GRAY
    doc.add_paragraph()
    cal = doc.add_paragraph("Calificación final: __________        Fecha: ____ / ____ / ______")
    cal.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_originality(doc):
    doc.add_heading("Declaración de originalidad", 1)
    doc.add_paragraph("Declaro que esta tesis es de mi autoría y constituye un trabajo original. Toda "
        "fuente, idea o resultado de terceros utilizado en su elaboración ha sido debidamente citado y "
        "referenciado conforme a las normas académicas vigentes. Declaro asimismo que este documento no "
        "ha sido presentado previamente para la obtención de otro grado o título.")
    for _ in range(3):
        doc.add_paragraph()
    f = doc.add_paragraph("__________________________________"); f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f.paragraph_format.space_after = Pt(0)
    a = doc.add_paragraph("Francisco Rietta"); a.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_simple(doc, titulo, cuerpo, italic=False, align=WD_ALIGN_PARAGRAPH.RIGHT):
    doc.add_heading(titulo, 1)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = align
    r = p.add_run(cuerpo); r.italic = italic; r.font.color.rgb = GRAY


def add_abstract_en(doc):
    doc.add_heading("Abstract", 1)
    doc.add_paragraph(
        "This thesis quantifies and compares the impact of global and national macroeconomic variables "
        "on the stock returns of copper firms with exposure to Chile over 2004–2024, contrasting how the "
        "shock is transmitted across the international and the Chilean equity markets. Using a triangulation "
        "design across three samples and a battery of time-series and panel-data models —unit-root tests, "
        "fixed-effects panels with Driscoll-Kraay standard errors, cointegration (ARDL, Johansen and "
        "Gregory-Hansen) and vector autoregressions with variance decomposition—, the study assesses the "
        "sensitivity of returns to macro-financial factors, its stability across copper-cycle phases, and "
        "its differential transmission across markets. The evidence supports the copper price as the main "
        "driver of returns and the dominance of global over local factors.")
    p = doc.add_paragraph(); p.add_run("Keywords: ").bold = True
    p.add_run("copper, stock returns, macroeconomic variables, Chile, panel data, cointegration, "
              "market segmentation.")


SIMBOLOS = [
    ("r_{i,t}", "Retorno logarítmico de la empresa i en el mes t"),
    ("β", "Coeficiente de sensibilidad (carga factorial) a un factor de riesgo"),
    ("α", "Velocidad de ajuste del mecanismo de corrección de error (VECM)"),
    ("Δ", "Operador de primera diferencia"),
    ("σ², σ", "Varianza y desviación estándar (volatilidad)"),
    ("ε_{i,t}", "Término de error idiosincrático"),
    ("λ", "Precio de mercado del riesgo asociado a un factor (APT)"),
    ("I(d)", "Serie integrada de orden d"),
    ("VR(q)", "Razón de varianzas a horizonte q (Lo-MacKinlay)"),
    ("γ", "Parámetro de asimetría (efecto apalancamiento) en GJR-GARCH"),
    ("ω_{i←j}(h)", "Fracción de la varianza de i explicada por el shock j a horizonte h (FEVD)"),
]


def add_simbolos(doc):
    doc.add_heading("Lista de símbolos", 1)
    doc.add_paragraph("Se resume a continuación la notación matemática empleada a lo largo del "
                      "documento.").paragraph_format.space_after = Pt(10)
    tbl = doc.add_table(rows=len(SIMBOLOS), cols=2); tbl.autofit = True
    for i, (sym, desc) in enumerate(SIMBOLOS):
        c0, c1 = tbl.rows[i].cells
        r = c0.paragraphs[0].add_run(sym); r.italic = True; r.font.name = "Cambria Math"
        c1.paragraphs[0].add_run(desc)
        for c in (c0, c1):
            _cell_pad(c)
            for rr in c.paragraphs[0].runs:
                rr.font.size = Pt(11)


def add_abbrev(doc):
    doc.add_heading("Lista de abreviaturas y siglas", 1)
    tbl = doc.add_table(rows=len(ABREV), cols=2); tbl.autofit = True
    for i, (sig, desc) in enumerate(ABREV):
        c0, c1 = tbl.rows[i].cells
        c0.paragraphs[0].add_run(sig).bold = True
        c1.paragraphs[0].add_run(desc)
        for c in (c0, c1):
            _cell_pad(c)
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(11); r.font.name = BODY_FONT


def _resumen_blocks():
    return [
        "Esta investigación cuantifica y compara el impacto de las variables macroeconómicas globales y "
        "nacionales sobre los retornos accionarios de las empresas de cobre con exposición a Chile durante "
        "2004–2024, contrastando su transmisión entre el mercado bursátil internacional y el chileno. "
        "Mediante un diseño de triangulación por tres muestras y un conjunto de modelos econométricos de "
        "series de tiempo y datos de panel —pruebas de raíz unitaria, panel de efectos fijos con errores "
        "Driscoll-Kraay, cointegración (ARDL, Johansen y Gregory-Hansen) y vectores autorregresivos con "
        "descomposición de varianza—, se evalúa la sensibilidad de los retornos a los factores "
        "macro-financieros, su estabilidad a lo largo de las fases del ciclo del cobre y la transmisión "
        "diferenciada entre mercados. La evidencia respalda al precio del cobre como principal determinante "
        "de los retornos y la dominancia de los factores globales sobre los locales.",
        "**Palabras clave:** cobre, retornos accionarios, variables macroeconómicas, Chile, datos de panel, "
        "cointegración, segmentación de mercados."]


# ---------- main ----------
def build():
    doc = Document()
    set_base_styles(doc)
    enable_hyphenation(doc)
    sec0 = doc.sections[0]
    sec0.left_margin = Cm(3.5); sec0.right_margin = Cm(3)
    sec0.top_margin = Cm(3); sec0.bottom_margin = Cm(3)
    sec0.different_first_page_header_footer = True  # portada sin encabezado/pie

    # Preliminares (romano)
    add_cover(doc); page(doc)
    add_approval(doc); page(doc)
    add_originality(doc); page(doc)
    add_simple(doc, "Dedicatoria", "[Espacio reservado para la dedicatoria.]", italic=True); page(doc)
    add_simple(doc, "Agradecimientos", "[Espacio reservado para los agradecimientos.]",
               align=WD_ALIGN_PARAGRAPH.JUSTIFY); page(doc)
    doc.add_heading("Resumen", 1)
    for blk in _resumen_blocks():
        add_runs_with_bold(doc.add_paragraph(), blk)
    page(doc)
    add_abstract_en(doc); page(doc)
    doc.add_heading("Índice general", 1); field_block(doc, 'TOC \\o "1-3" \\h \\z \\u', "Actualizar campo (clic derecho → Actualizar campos)."); page(doc)
    doc.add_heading("Índice de tablas", 1); field_block(doc, 'TOC \\h \\z \\c "Tabla"', "Actualizar campo."); page(doc)
    doc.add_heading("Índice de figuras", 1); field_block(doc, 'TOC \\h \\z \\c "Figura"', "Actualizar campo."); page(doc)
    add_abbrev(doc); page(doc)
    add_simbolos(doc)
    set_pgnum_format(sec0, "lowerRoman", start=1)
    running_header(sec0)

    # Cuerpo (arábigo)
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.left_margin = Cm(3.5); body.right_margin = Cm(3)
    body.top_margin = Cm(3); body.bottom_margin = Cm(3)
    body.different_first_page_header_footer = False
    set_pgnum_format(body, "decimal", start=1)
    running_header(body)
    for k, fname in enumerate(CAPS):
        render_markdown(doc, (DOCS / fname).read_text(encoding="utf-8"),
                        hanging=(fname == "referencias.md"))
        if k < len(CAPS) - 1:
            page(doc)

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"[ok] {OUT}  ({OUT.stat().st_size//1024} KB)")
    return OUT


if __name__ == "__main__":
    build()
