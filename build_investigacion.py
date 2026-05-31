"""
build_investigacion.py
======================
Convierte TODOS los documentos de investigación (carpeta Investigacion/) a Word (.docx)
con un DISEÑO ESTILO APPLE PROFESIONAL y, además, un Compendio único que los reúne.

Principios de diseño (Apple + sugerencias de la MIT *Specifications for Thesis Preparation*):
- Tipografía sans serif (SF Pro Display) — la MIT sugiere sans serif por legibilidad.
- Mucho espacio en blanco, márgenes amplios, texto alineado a la izquierda (sin justificar).
- Jerarquía tipográfica clara: títulos grandes, "kicker" de color, texto secundario gris.
- Color de texto Apple #1D1D1F; gris secundario #6E6E73; acento azul #0071E3.
- Índice (la MIT lo recomienda), paginación consecutiva, enlaces clicables (MIT).

Salida: outputs/investigacion/*.docx  +  outputs/investigacion/Compendio_Investigacion.docx
"""
from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
INV = ROOT / "Investigacion"
OUTDIR = ROOT / "outputs" / "investigacion"

# ---- paleta Apple ----
INK = RGBColor(0x1D, 0x1D, 0x1F)     # texto principal (Apple)
GRAY = RGBColor(0x6E, 0x6E, 0x73)    # texto secundario (Apple)
ACCENT = RGBColor(0x00, 0x71, 0xE3)  # azul Apple
HAIR = "D2D2D7"                       # filete fino (Apple separators)
# Tipografía: Segoe UI (sans serif limpia, presente en Windows y de aspecto Apple-like).
# Se descartó "SF Pro Display" por substitución a una variante oblicua en la conversión a PDF.
DISPLAY = "Segoe UI Semibold"
BODY = "Segoe UI"
MONO = "Consolas"

# documentos de investigación (ruta relativa a Investigacion/, título amigable, subtítulo)
DOCMAP = [
    ("01_estado_del_arte/estado_del_arte.md", "Estado del Arte",
     "Qué se sabe, qué vacío llena la tesis y cómo se posiciona en la literatura"),
    ("02_historia_del_cobre/historia_del_cobre.md", "Historia del Cobre",
     "Historia económica del cobre y de la minería en Chile, 2004–2024"),
    ("03_literatura/revision_literatura.md", "Revisión de Literatura",
     "Marco de antecedentes estructurado"),
    ("03_literatura/lit_chilena.md", "Literatura Chilena y Latinoamericana",
     "Tipo de cambio, cobre y mercado de capitales local"),
    ("03_literatura/lit_emergentes.md", "Literatura de Mercados Emergentes",
     "Commodities y mercados emergentes"),
    ("03_literatura/busqueda_literatura.md", "Estrategia de Búsqueda Bibliográfica",
     "Términos, fuentes y criterios de inclusión"),
    ("04_metodologia_y_mejoras/mejoras_metodologicas.md", "Mejoras Metodológicas",
     "Batería econométrica de segunda generación"),
    ("04_metodologia_y_mejoras/diseno_triangulacion.md", "Diseño de Triangulación",
     "Tres muestras: cobre internacional, Pucobre y sector minero chileno"),
    ("04_metodologia_y_mejoras/matriz_consistencia.md", "Matriz de Consistencia",
     "Problema – objetivos – hipótesis – variables – método"),
    ("05_referencias/referencias.md", "Referencias",
     "Listado completo en formato APA 7"),
    ("06_datos_y_hallazgos/hallazgos_datos.md", "Hallazgos de los Datos",
     "Resultados de las series descargadas (FRED + Yahoo Finance)"),
    ("06_datos_y_hallazgos/guia_descarga_bcch.md", "Guía de Descarga (BCCh)",
     "Descarga del EMBI desde el Banco Central de Chile"),
    ("06_datos_y_hallazgos/reproducibilidad.md", "Reproducibilidad",
     "Flujo de trabajo, entorno y pasos para replicar"),
]


# ---------- estilos ----------
def set_styles(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY; st.font.size = Pt(11.5); st.font.color.rgb = INK
    pf = st.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(10)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT   # Apple: ragged-right, sin justificar
    pf.widow_control = True
    heads = [("Heading 1", 22, INK), ("Heading 2", 16, INK), ("Heading 3", 13, ACCENT)]
    for h, sz, col in heads:
        s = doc.styles[h]
        # Segoe UI Semibold aporta el peso; bold=False evita el faux-bold sobre semibold.
        s.font.name = DISPLAY; s.font.size = Pt(sz); s.font.bold = False; s.font.italic = False
        s.font.color.rgb = col
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        s.paragraph_format.line_spacing = 1.1
        s.paragraph_format.space_before = Pt(18 if h != "Heading 1" else 10)
        s.paragraph_format.space_after = Pt(8 if h == "Heading 1" else 4)


# ---------- utilidades OOXML ----------
def _fld(run, instr, default="1"):
    for t, attr in [("begin", None), (None, instr), ("separate", None), ("text", default), ("end", None)]:
        if attr is not None:
            it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = attr
            run._r.append(it)
        elif t == "text":
            tt = OxmlElement("w:t"); tt.text = default; run._r.append(tt)
        else:
            fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), t); run._r.append(fc)


def page_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.text = ""
    run = p.add_run(); _fld(run, "PAGE")
    run.font.size = Pt(9); run.font.name = BODY; run.font.color.rgb = GRAY


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                          is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts"); rfonts.set(qn("w:ascii"), BODY); rfonts.set(qn("w:hAnsi"), BODY)
    rPr.append(rfonts)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0071E3"); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; new_run.append(t)
    hl.append(new_run); paragraph._p.append(hl)


def hairline(doc, space_before=4, space_after=10, color=HAIR, sz=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); pPr.append(pbdr)
    b = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", str(sz)), ("w:space", "1"), ("w:color", color)):
        b.set(qn(k), v)
    pbdr.append(b)
    return p


# ---------- runs con formato + enlaces clicables (MIT) ----------
URL_RE = r"https?://[^\s)\]]+"


def add_runs(paragraph, text):
    # negrita, código, matemática, itálica, enlaces markdown [t](u), URLs sueltas
    pat = (r"(\[[^\]]+\]\((?:https?://|www\.)[^)]+\)|\*\*.+?\*\*|`[^`]+`|"
           r"\$[^$\n]+\$|\*[^*\n]+?\*|" + URL_RE + r")")
    for part in re.split(pat, text):
        if not part:
            continue
        m = re.match(r"\[([^\]]+)\]\(((?:https?://|www\.)[^)]+)\)", part)
        if m:
            url = m.group(2)
            add_hyperlink(paragraph, url if url.startswith("http") else "http://" + url, m.group(1))
        elif re.fullmatch(URL_RE, part):
            add_hyperlink(paragraph, part, part)
        elif part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = MONO; r.font.size = Pt(10)
        elif part.startswith("$") and part.endswith("$"):
            paragraph.add_run(part[1:-1]).italic = True
        elif len(part) > 2 and part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def clean_heading(text):
    return re.sub(r"\s*\((borrador|texto)[^)]*\)", "", text).strip()


# ---------- tablas estilo Apple (líneas finas, encabezado con fondo claro) ----------
def _cell_border(cell, edge, sz, color):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders"); tcPr.append(borders)
    el = borders.find(qn("w:" + edge))
    if el is None:
        el = OxmlElement("w:" + edge); borders.append(el)
    for k, v in (("w:val", "single"), ("w:sz", str(sz)), ("w:space", "0"), ("w:color", color)):
        el.set(qn(k), v)


def _cell_pad(cell):
    tcPr = cell._tc.get_or_add_tcPr(); mar = OxmlElement("w:tcMar")
    for edge, w in (("top", 80), ("bottom", 80), ("left", 110), ("right", 110)):
        e = OxmlElement("w:" + edge); e.set(qn("w:w"), str(w)); e.set(qn("w:type"), "dxa"); mar.append(e)
    tcPr.append(mar)


def _cell_shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:color"), "auto"); sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not cells:
        return
    tbl = doc.add_table(rows=len(cells), cols=len(cells[0]))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT; tbl.autofit = True
    nrows = len(cells)
    for i, row in enumerate(cells):
        for j, val in enumerate(row):
            if j >= len(tbl.rows[i].cells):
                continue
            c = tbl.rows[i].cells[j]; cp = c.paragraphs[0]; cp.text = ""
            cp.paragraph_format.space_after = Pt(2); cp.paragraph_format.line_spacing = 1.05
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_runs(cp, val.replace("**", "") if i == 0 else val)
            for run in cp.runs:
                run.font.size = Pt(10); run.font.name = BODY
                if i == 0:
                    run.bold = True; run.font.color.rgb = INK
            _cell_pad(c)
            if i == 0:
                _cell_shade(c, "F5F5F7")          # gris Apple muy claro en cabecera
                _cell_border(c, "bottom", 6, HAIR)
            else:
                _cell_border(c, "bottom", 4, HAIR)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc, cap_text, path):
    img = ROOT / path
    if not img.exists():
        return
    doc.add_picture(str(img), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
    if cap_text:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(cap_text); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GRAY


def unwrap_paragraphs(md):
    def is_break(s):
        return (not s.strip()) or bool(re.match(r"\s*(#{1,6}\s|\||!\[|>|---|\$\$)", s))
    def is_cont(s):
        return bool(s.strip()) and not re.match(r"\s*(#{1,6}\s|[-*]\s|\d+\.\s|\||!\[|>|---|\$\$)", s)
    out = []
    for raw in md.split("\n"):
        s = raw.rstrip()
        if out and is_cont(s) and not is_break(out[-1]):
            out[-1] += " " + s.strip()
        else:
            out.append(s)
    return "\n".join(out)


def render(doc, md_text, demote=0, skip_h1=False):
    """demote: baja N niveles los encabezados (para el compendio). skip_h1: omite el primer H1."""
    lines = unwrap_paragraphs(md_text).split("\n"); i = 0; tablebuf = []; seen_h1 = False
    def flush():
        nonlocal tablebuf
        if tablebuf:
            add_table(doc, tablebuf); tablebuf = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("|") and "|" in line[1:]:
            tablebuf.append(line); i += 1; continue
        flush()
        if not line.strip() or line.startswith(">") or line.startswith("---"):
            i += 1; continue
        m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if m:
            add_figure(doc, m.group(1), m.group(2)); i += 1; continue
        meq = re.match(r"^\s*\$\$(.+?)\$\$\s*$", line)
        if meq:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(meq.group(1).strip()); r.italic = True
            i += 1; continue
        hm = re.match(r"^(#{1,6})\s+(.*)", line)
        if hm:
            lvl = len(hm.group(1))
            if lvl == 1 and skip_h1 and not seen_h1:
                seen_h1 = True; i += 1; continue
            seen_h1 = True
            lvl = min(3, lvl + demote)
            doc.add_heading(clean_heading(hm.group(2)), lvl); i += 1; continue
        if re.match(r"^\s*[-*] ", line):
            add_runs(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*] ", "", line)); i += 1; continue
        mn = re.match(r"^\s*(\d+)\.\s+(.*)", line)
        if mn:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8); p.paragraph_format.first_line_indent = Cm(-0.8)
            add_runs(p, f"{mn.group(1)}. {mn.group(2)}"); i += 1; continue
        add_runs(doc.add_paragraph(), line); i += 1
    flush()


# ---------- portada estilo Apple ----------
def _spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)


def apple_cover(doc, kicker, title, subtitle, meta_lines):
    _spacer(doc, 4)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(kicker.upper()); r.font.name = DISPLAY; r.font.size = Pt(12)
    r.font.color.rgb = ACCENT
    # espaciado entre letras del kicker
    rPr = r._r.get_or_add_rPr(); sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), "40"); rPr.append(sp)
    pt = doc.add_paragraph(); pt.paragraph_format.space_after = Pt(8)
    rt = pt.add_run(title); rt.font.name = DISPLAY; rt.font.size = Pt(34)
    rt.font.color.rgb = INK
    pt.paragraph_format.line_spacing = 1.02
    if subtitle:
        ps = doc.add_paragraph(); ps.paragraph_format.space_after = Pt(0)
        rs = ps.add_run(subtitle); rs.font.name = BODY; rs.font.size = Pt(15); rs.font.color.rgb = GRAY
        ps.paragraph_format.line_spacing = 1.2
    _spacer(doc, 1)
    hairline(doc, space_before=8, space_after=14)
    for label, val in meta_lines:
        pm = doc.add_paragraph(); pm.paragraph_format.space_after = Pt(2)
        rl = pm.add_run(label + "   "); rl.font.name = DISPLAY; rl.font.size = Pt(10.5)
        rl.font.color.rgb = INK
        rv = pm.add_run(val); rv.font.name = BODY; rv.font.size = Pt(10.5); rv.font.color.rgb = GRAY


def setup_section(section):
    section.left_margin = Cm(3.2); section.right_margin = Cm(3.2)
    section.top_margin = Cm(3.0); section.bottom_margin = Cm(2.6)


# ---------- builders ----------
def build_one(rel, title, subtitle):
    md = (INV / rel).read_text(encoding="utf-8")
    doc = Document(); set_styles(doc)
    sec = doc.sections[0]; setup_section(sec)
    sec.different_first_page_header_footer = True
    apple_cover(doc, "Investigación · Tesis F. Rietta", title, subtitle, [
        ("Proyecto", "Macroeconomía y retornos del sector cobre en Chile (2004–2024)"),
        ("Documento", title),
        ("Formato", "Estilo Apple · sugerencias MIT (sans serif, índice, enlaces)"),
    ])
    doc.add_page_break()
    page_footer(sec)
    render(doc, md, demote=0, skip_h1=False)
    OUTDIR.mkdir(parents=True, exist_ok=True)
    out = OUTDIR / (Path(rel).stem + ".docx")
    doc.save(out)
    return out


def build_compendium():
    doc = Document(); set_styles(doc)
    sec = doc.sections[0]; setup_section(sec)
    sec.different_first_page_header_footer = True
    apple_cover(doc, "Compendio de Investigación", "Compendio de Investigación",
                "Estado del arte, historia del cobre, literatura, metodología, referencias y datos",
                [("Proyecto", "Macroeconomía y retornos del sector cobre en Chile (2004–2024)"),
                 ("Autor", "Francisco Rietta"),
                 ("Documentos", f"{len(DOCMAP)} secciones de investigación"),
                 ("Formato", "Estilo Apple · sugerencias MIT")])
    doc.add_page_break()
    page_footer(sec)
    # Índice (MIT: strongly recommended)
    h = doc.add_heading("Índice", 1)
    pf = doc.add_paragraph(); r = pf.add_run("Actualice el campo con clic derecho → Actualizar campos.")
    r.font.color.rgb = GRAY; r.italic = True; r.font.size = Pt(10)
    p = doc.add_paragraph(); run = p.add_run(); _fld(run, 'TOC \\o "1-2" \\h \\z \\u', "Índice")
    doc.add_page_break()
    # cuerpo: cada documento como capítulo
    for k, (rel, title, subtitle) in enumerate(DOCMAP):
        doc.add_heading(title, 1)
        ps = doc.add_paragraph(); rs = ps.add_run(subtitle); rs.font.color.rgb = GRAY; rs.font.size = Pt(12)
        hairline(doc, space_before=2, space_after=10)
        render(doc, (INV / rel).read_text(encoding="utf-8"), demote=1, skip_h1=True)
        if k < len(DOCMAP) - 1:
            doc.add_page_break()
    OUTDIR.mkdir(parents=True, exist_ok=True)
    out = OUTDIR / "Compendio_Investigacion.docx"
    doc.save(out)
    return out


def build():
    outs = []
    for rel, title, subtitle in DOCMAP:
        outs.append(build_one(rel, title, subtitle))
        print(f"[ok] {outs[-1].name}")
    comp = build_compendium()
    print(f"[ok] {comp.name}  (compendio)")
    outs.append(comp)
    return outs


if __name__ == "__main__":
    build()
